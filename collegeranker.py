import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document

# Title of the app
st.title("College Ranking Analysis Tool")

# File uploader
uploaded_file = st.file_uploader("Upload your data file (CSV format)", type=["csv"])

if uploaded_file:
    # Load the data
    data = pd.read_csv(uploaded_file)
    st.write("### Data Overview")
    st.write("Below is a preview of the uploaded data:")
    st.write(data.head())

    # Ensure consistent column names (strip spaces)
    data.columns = data.columns.str.strip()

    # Pivot data to create year-wise ranking columns
    pivot_data = data.pivot_table(
        index=["College Name", "City_name", "Stream", "Agency Name", "college_type"],
        columns="year",
        values="end_ranking_of_college",
        aggfunc="first"
    ).reset_index()
    pivot_data.columns = [str(col) if isinstance(col, int) else col for col in pivot_data.columns]

    st.write("### Pivoted Data with Year-Wise Rankings")
    st.write("The table below shows rankings split by year:")
    st.write(pivot_data)

    # Prepare a Word document for output
    document = Document()
    document.add_heading("College Ranking Analysis", level=1)

    # Filter options
    city = st.selectbox("Select City", ["All"] + list(data["City_name"].unique()))
    stream = st.selectbox("Select Stream", ["All"] + list(data["Stream"].unique()))
    agency = st.selectbox("Select Agency", ["All"] + list(data["Agency Name"].unique()))
    college_type = st.selectbox("Select College Type", ["All"] + list(data["college_type"].unique()))

    # Generate a keyword-rich phrase
    keyword = f"Top Colleges in {city if city != 'All' else 'India'} for {stream if stream != 'All' else 'various streams'}"
    if agency != "All":
        keyword += f", Ranked by {agency}"
    if college_type != "All":
        keyword += f" ({college_type.capitalize()} Colleges)"
    st.write(f"### Optimized Content for: {keyword}")
    document.add_heading(keyword, level=2)

    # Generate SEO-focused content
    seo_text = f"""
    {keyword} provide an excellent overview of the best institutes in the country. These colleges are known for their academic excellence, placement records, and overall contribution to industry and research. 
    With rankings provided by {agency if agency != 'All' else 'multiple agencies'}, students can make informed decisions about their education. Explore year-wise rankings to understand how these colleges have performed over time.
    """
    st.write(seo_text)
    document.add_paragraph(seo_text)

    # Apply filters to the pivoted data
    filtered_data = pivot_data.copy()
    if city != "All":
        filtered_data = filtered_data[filtered_data["City_name"] == city]
    if stream != "All":
        filtered_data = filtered_data[filtered_data["Stream"] == stream]
    if agency != "All":
        filtered_data = filtered_data[filtered_data["Agency Name"] == agency]
    if college_type != "All":
        filtered_data = filtered_data[filtered_data["college_type"] == college_type]

    # Display filtered data
    st.write("### Filtered Data")
    st.write("The table below shows filtered rankings:")
    st.write(filtered_data)

    # Add filtered table to Word document
    document.add_paragraph(f"Filtered Data for City: {city}, Stream: {stream}, Agency: {agency}")
    table = document.add_table(rows=1, cols=len(filtered_data.columns))
    for i, column in enumerate(filtered_data.columns):
        table.rows[0].cells[i].text = column
    for _, row in filtered_data.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    # Visualizations
    st.write("### Visualizations")
    st.write("Select one or more colleges to visualize their year-wise rankings.")
    colleges = st.multiselect("Select Colleges", filtered_data["College Name"].unique())

    if colleges:
        fig, ax = plt.subplots()
        for college in colleges:
            college_data = filtered_data[filtered_data["College Name"] == college]
            years = [col for col in college_data.columns if col.isdigit()]
            if not years:
                continue
            rankings = college_data[years].iloc[0]
            ax.plot(years, rankings, marker="o", label=college)

        ax.set_title(f"Year-Wise Rankings for Selected Colleges: {keyword}")
        ax.set_ylabel("Ranking")
        ax.set_xlabel("Year")
        ax.legend()
        st.pyplot(fig)

        # Save graph to Word document
        fig.savefig("rankings_graph.png")
        document.add_picture("rankings_graph.png")
        document.add_paragraph("Graph: Year-wise rankings for selected colleges")

    # Add download option for the Word file
    st.write("### Download Results")
    st.write("Download all tables, graphs, and SEO-focused content in a Word document.")
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    st.download_button(
        label="Download Word File",
        data=buffer,
        file_name="college_ranking_analysis.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
