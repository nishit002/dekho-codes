import streamlit as st
import pandas as pd
import docx
from io import BytesIO
import re
import base64
from textblob import TextBlob

# Function to convert text to proper case, ignoring abbreviations
def proper_case_except_abbreviations(text):
    def capitalize_word(word):
        if word.isupper():  # Keep abbreviations in uppercase
            return word
        else:
            return word.capitalize()  # Capitalize only non-abbreviation words
    return ' '.join([capitalize_word(word) for word in re.split(r'(\W+)', text)])

# Function to correct spelling mistakes and clean text using TextBlob
def correct_text(text):
    if isinstance(text, str):
        text = TextBlob(text).correct()  # Apply spelling correction
        text = proper_case_except_abbreviations(str(text))  # Convert to proper case
    return text

# Function to load the predefined Excel file from the local path or URL
@st.cache_data
def load_predefined_excel():
    # Load the Excel file from a local path or URL (your GitHub file path)
    file_url = "https://raw.githubusercontent.com/your-github-username/dekho-codes/main/ALL_CCM_ALL_Template.xlsx"
    df = pd.read_excel(file_url, header=0)
    
    # College Name is in the first column (Column 0)
    df['College Name'] = df.iloc[:, 0].astype(str)
    
    # Drop rows where College Name is NaN
    df = df.dropna(subset=['College Name'])
    
    # Count the number of filled (non-empty) fields for each college
    df['Filled Fields Count'] = df.notna().sum(axis=1) - 1  # Subtract 1 to exclude the college name column
    
    return df

# Function to create a Word document with a transposed table (with borders)
def create_transposed_word_table(college_data):
    doc = docx.Document()
    doc.add_heading('College Details', 0)
    
    table = doc.add_table(rows=len(college_data.columns), cols=2)  # Transpose: Rows = number of columns in original data
    table.style = 'Table Grid'  # Add borders to the table
    
    # Add headers and values in a transposed manner
    for i, col in enumerate(college_data.columns):
        table.rows[i].cells[0].text = col  # Column name
        table.rows[i].cells[1].text = str(college_data[col].values[0])  # Corresponding value
    
    # Create a BytesIO object and save the document into it
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# Function to create a transposed HTML table
def create_transposed_html_table(college_data):
    transposed_df = college_data.T  # Transpose the dataframe
    html_table = transposed_df.to_html(header=False, border=1)
    return html_table

# Streamlit app
st.title('College Information Table Generator with On-Demand Text Correction')

# Load the predefined Excel file automatically
st.write("Loading predefined Excel file...")
processed_df = load_predefined_excel()

# Create a display list combining college name and the filled field count
processed_df['College Display'] = processed_df['College Name'] + " (Fields Filled: " + processed_df['Filled Fields Count'].astype(str) + ")"

# Let the user manually select the college from the formatted "College Display" column
college_list = processed_df['College Display'].tolist()
selected_college = st.selectbox("Select a college", college_list)

# Filter the data for the selected college
if selected_college:
    st.write(f"Details for {selected_college}:")
    
    # Get the college row for the selected college name
    college_data = processed_df[processed_df['College Display'] == selected_college].drop(columns=['College Display', 'Filled Fields Count'])
    
    # Filter out empty columns (only show columns where data is available)
    college_data = college_data.loc[:, college_data.notna().any()]
    
    # Display the original data in transposed format
    st.dataframe(college_data.T)
    
    # Add a button to correct spelling and grammar
    if st.button('Correct Spelling and Grammar'):
        # Apply text correction only to the displayed data
        college_data = college_data.applymap(lambda x: correct_text(x) if isinstance(x, str) else x)
        
        # Display corrected data
        st.write("Corrected Data:")
        st.dataframe(college_data.T)
    
    # Option to download the transposed and corrected college data as a Word file
    buffer_word = create_transposed_word_table(college_data)
    st.download_button(label="Download College Data as Word", data=buffer_word, file_name=f"{selected_college}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    # Option to download the transposed and corrected college data as HTML
    html_table = create_transposed_html_table(college_data)
    b64_html = base64.b64encode(html_table.encode()).decode()  # Encode HTML to base64
    href_html = f'<a href="data:text/html;base64,{b64_html}" download="{selected_college}.html">Download College Data as HTML</a>'
    st.markdown(href_html, unsafe_allow_html=True)
    
    # Option to copy the transposed table to clipboard (HTML format)
    st.text_area("HTML Table (Copy this):", value=html_table, height=200)
