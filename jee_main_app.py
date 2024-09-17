import streamlit as st
import pandas as pd
import plotly.express as px
from io import BytesIO
from docx import Document
from docx.shared import Inches

# Function to merge the two dataframes based on common key columns
def merge_dataframes(df1, df2, key_columns):
    return pd.merge(df1, df2, on=key_columns, suffixes=('_Year1', '_Year2'))

# Function to standardize the dataframe column names
def preprocess_dataframe(df):
    df.rename(columns={
        'Institute': 'College Name',
        'Academic Program Name': 'Course Name',
        'Quota': 'Quota',
        'Seat Type': 'Seat Type',
        'Gender': 'Gender',
        'Opening Rank': 'Opening Rank',
        'Closing Rank': 'Closing Rank'
    }, inplace=True)
    return df

# Function to create a Word document with table and chart
def create_word_file(df, fig, filename='Comparison_Report.docx'):
    doc = Document()
    doc.add_heading('Comparison Report', level=1)
    doc.add_paragraph('Comparison of Opening and Closing Ranks:')
    
    # Add the DataFrame as a table in the Word document
    t = doc.add_table(df.shape[0]+1, df.shape[1])
    for j, column in enumerate(df.columns):
        t.cell(0, j).text = column
        for i, val in enumerate(df[column].values):
            t.cell(i+1, j).text = str(val)
    
    # Add the chart to the Word document
    doc.add_paragraph('Rank Changes Over Years:')
    fig_bytes = BytesIO()
    fig.write_image(fig_bytes, format='png')
    fig_bytes.seek(0)
    doc.add_picture(fig_bytes, width=Inches(6))
    
    doc.save(filename)
    return filename

# Function to create a Plotly chart
def create_plotly_chart(df, key_columns):
    fig = px.line(df, x='Course Name', y=['Opening Rank_Year1', 'Opening Rank_Year2'],
                  color='Quota', line_dash='Gender', 
                  hover_name='Course Name', markers=True,
                  title='Rank Changes Over Years')
    fig.update_layout(legend_title_text='Quota-Gender')
    fig.update_xaxes(title_text='Course Name')
    fig.update_yaxes(title_text='Rank')
    return fig

# Main Streamlit application starts here
st.title('College Rank Comparison Tool')

# Input for years and file uploads
year1 = st.text_input("Enter year for the first file:")
year2 = st.text_input("Enter year for the second file:")
file1 = st.file_uploader(f"Upload the Excel file for {year1}", type=["xlsx"], key="file1")
file2 = st.file_uploader(f"Upload the Excel file for {year2}", type=["xlsx"], key="file2")

if file1 and file2 and year1 and year2:
    df1 = pd.read_excel(file1)
    df2 = pd.read_excel(file2)
    
    # Preprocess the dataframes to standardize column names
    df1 = preprocess_dataframe(df1)
    df2 = preprocess_dataframe(df2)
    
    # Dropdown to select college
    college_names = pd.concat([df1['College Name'], df2['College Name']]).unique()
    selected_college = st.selectbox('Select a College', college_names)
    
    if selected_college:
        # Optional dropdown to select a course
        course_names = pd.concat([df1[df1['College Name'] == selected_college]['Course Name'],
                                  df2[df2['College Name'] == selected_college]['Course Name']]).unique()
        selected_course = st.selectbox('Select a Course (optional)', ['Any'] + list(course_names))
        
        if selected_course and selected_course != 'Any':
            df1 = df1[(df1['College Name'] == selected_college) & (df1['Course Name'] == selected_course)]
            df2 = df2[(df2['College Name'] == selected_college) & (df2['Course Name'] == selected_course)]
        else:
            df1 = df1[df1['College Name'] == selected_college]
            df2 = df2[df2['College Name'] == selected_college]
        
        # Dropdowns for Quota, Gender, and Seat Type
        selected_quota = st.selectbox('Select Quota', df1['Quota'].unique())
        selected_gender = st.selectbox('Select Gender', df1['Gender'].unique())
        selected_seat_type = st.selectbox('Select Seat Type', df1['Seat Type'].unique())
        
        # Filter data based on selected Quota, Gender, and Seat Type
        filtered_df = df1[(df1['Quota'] == selected_quota) &
                          (df1['Gender'] == selected_gender) &
                          (df1['Seat Type'] == selected_seat_type)]
        
        # Merge the filtered data with the second dataset
        key_columns = ['College Name', 'Course Name', 'Quota', 'Seat Type', 'Gender']
        merged_df = merge_dataframes(filtered_df, df2, key_columns)
        
        if not merged_df.empty:
            # Display the merged table
            st.write("Comparison of Opening and Closing Ranks for Selected Filters:")
            st.table(merged_df[['Course Name', 'Quota', 'Seat Type', 'Gender', 
                                'Opening Rank_Year1', 'Closing Rank_Year1', 
                                'Opening Rank_Year2', 'Closing Rank_Year2']])
            
            # Create and display the Plotly chart
            fig = create_plotly_chart(merged_df, key_columns)
            st.plotly_chart(fig)
            
            # Provide a download button for the Word report
            doc_filename = create_word_file(merged_df, fig)
            with open(doc_filename, "rb") as file:
                st.download_button(
                    label="Download Report as Word Document",
                    data=file,
                    file_name=doc_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.write("No matching data found for the selected filters.")
