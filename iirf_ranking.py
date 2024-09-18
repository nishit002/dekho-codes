import streamlit as st
import pandas as pd
import plotly.express as px
from googlesearch import search
import requests
from bs4 import BeautifulSoup
from io import BytesIO
from docx import Document
from docx.shared import Inches
import os

# Extract the ranking name from the file name
def extract_ranking_name(filename):
    return os.path.splitext(filename)[0]

# Process Excel files and return all combined data
def process_excel_files(uploaded_files, ranking_names):
    all_data = pd.DataFrame()
    for i, uploaded_file in enumerate(uploaded_files):
        df = pd.read_excel(uploaded_file, usecols="A:D", names=['Rank', 'College Name', 'City', 'State'])
        df['Rank'] = df['Rank'].apply(lambda x: f"#{int(x)}" if pd.notna(x) and isinstance(x, (int, float)) else 'N/A')
        df['Ranking Stream'] = ranking_names[i]
        all_data = pd.concat([all_data, df], ignore_index=True)
    return all_data

# Perform Google search for the IIRF 2023 ranking
def google_search_college_ranking(college_name):
    query = f"{college_name} IIRF Ranking 2023"
    result_text = ""
    try:
        search_results = search(query, stop=1, pause=2)
        for result in search_results:
            page = requests.get(result)
            soup = BeautifulSoup(page.content, 'html.parser')
            result_text = soup.get_text()
            break
    except Exception as e:
        print(f"Error fetching data from Google: {e}")
    return result_text

# Generate the Word document report
def create_word_report(college_name, paragraph, table_data, graph_bytes):
    doc = Document()
    doc.add_heading(f"{college_name} Ranking Report", level=0)

    # Add the summary paragraph
    doc.add_paragraph(paragraph)

    # Add the table to the Word document
    if not table_data.empty:
        table = doc.add_table(rows=1, cols=len(table_data.columns))
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(table_data.columns):
            hdr_cells[i].text = str(col_name)
        for index, row in table_data.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

    # Add the graph to the Word document
    doc.add_paragraph('Ranking Graph:')
    doc.add_picture(graph_bytes, width=Inches(6))

    # Save the document to a BytesIO object to return
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

# Generate the ranking paragraph based on college data and Google search results
def generate_ranking_paragraph(college_name, college_data, google_ranking):
    files_found_in = len(college_data['Ranking Stream'].unique())
    paragraph = f"{college_name} has been ranked in {files_found_in} different ranking streams.\n"
    for ranking_stream, group in college_data.groupby('Ranking Stream'):
        file_ranking = group.iloc[0]['Rank']
        paragraph += f"In the '{ranking_stream}' ranking stream, {college_name} has a rank of {file_ranking}.\n"
    if google_ranking:
        if 'Rank' in google_ranking:
            paragraph += f"According to the 2023 IIRF ranking from Google, {college_name} was ranked at {google_ranking}. "
        current_rank = int(college_data['Rank'].str.replace('#', '').iloc[0]) if pd.notna(college_data['Rank'].iloc[0]) else None
        previous_rank = int(google_ranking.split()[1]) if google_ranking else None
        if current_rank and previous_rank:
            if current_rank < previous_rank:
                paragraph += f"The rank has improved from {previous_rank} to {current_rank} in this year's rankings."
            elif current_rank > previous_rank:
                paragraph += f"The rank has declined from {previous_rank} to {current_rank} in this year's rankings."
            else:
                paragraph += f"The rank remains unchanged at {current_rank} compared to last year."
    return paragraph

# Generate the graph for the selected college
def display_graph_for_college(college_data, college_name, watermark_text=""):
    fig = px.bar(
        college_data,
        x='Ranking Stream',
        y='Rank',
        color='Ranking Stream',
        labels={'Rank': 'Rank'},
        title=f'Rankings for {college_name} Across Multiple Streams',
    )
    fig.update_traces(texttemplate='%{y}', textposition='outside')
    fig.update_layout(uniformtext_minsize=8, uniformtext_mode='hide')
    fig.add_annotation(
        text=watermark_text,
        xref="paper", yref="paper",
        x=0.5, y=0.5, opacity=0.2,
        font=dict(size=40, color="lightgray"),
        showarrow=False
    )
    st.plotly_chart(fig)
    graph_bytes = BytesIO()
    fig.write_image(graph_bytes, format='png')
    graph_bytes.seek(0)
    return graph_bytes

# Main function for the Streamlit app
def main():
    st.title("College Ranking Analyzer")
    uploaded_files = st.file_uploader("Upload Excel files", type=["xlsx"], accept_multiple_files=True)

    if uploaded_files:
        ranking_names = [extract_ranking_name(file.name) for file in uploaded_files]
        all_data = process_excel_files(uploaded_files, ranking_names)

        ranking_type = st.selectbox("Select Ranking Type", [
            'College-wise', 'Stream+City-wise', 'Stream+State-wise', 'All Colleges in a City', 'All Colleges in a State'
        ])

        if ranking_type == 'College-wise':
            colleges = all_data['College Name'].dropna().unique()
            selected_college = st.selectbox('Select a College', colleges)
            if selected_college:
                college_data = all_data[all_data['College Name'] == selected_college]
                st.write(college_data[['Ranking Stream', 'Rank', 'City', 'State']].sort_values(by='Rank'))

                # Google Search for College Ranking Text
                google_ranking_text = google_search_college_ranking(selected_college)

                # Generate Summary Paragraph for the selected college
                summary_paragraph = generate_ranking_paragraph(selected_college, college_data, google_ranking_text)
                st.write(summary_paragraph)

                # Display Graph for the selected college
                watermark_text = "Your Watermark"
                graph_bytes = display_graph_for_college(college_data, selected_college, watermark_text)

                # Create and Download Word Report
                word_file = create_word_report(selected_college, summary_paragraph, college_data[['Ranking Stream', 'Rank', 'City', 'State']], graph_bytes)
                
                # Add a download button for the Word document
                st.download_button(
                    label="Download Report as Word Document",
                    data=word_file,
                    file_name=f"{selected_college}_Ranking_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        elif ranking_type == 'Stream+City-wise':
            ranking_streams = all_data['Ranking Stream'].dropna().unique()
            selected_stream = st.selectbox('Select Ranking Stream', ranking_streams)
            cities = all_data['City'].dropna().unique()
            selected_city = st.selectbox('Select City', cities)
            if selected_stream and selected_city:
                stream_city_data = all_data[(all_data['Ranking Stream'] == selected_stream) & (all_data['City'] == selected_city)]
                st.write(stream_city_data[['College Name', 'Rank', 'City', 'State']].sort_values(by='Rank'))

                # Generate a simple summary paragraph
                city_summary = f"Overview of all colleges in {selected_city} for stream '{selected_stream}'."
                st.write(city_summary)

                # Display Graph for colleges in the selected city and stream
                watermark_text = "Your Watermark"
                graph_bytes = display_graph_for_college(stream_city_data, selected_city, watermark_text)

                # Create and Download Word Report
                word_file = create_word_report(f"{selected_city} - {selected_stream}", city_summary, stream_city_data[['College Name', 'Rank', 'City', 'State']], graph_bytes)
                
                # Add a download button for the Word document
                st.download_button(
                    label="Download Report as Word Document",
                    data=word_file,
                    file_name=f"{selected_city}_{selected_stream}_Ranking_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        elif ranking_type == 'Stream+State-wise':
            ranking_streams = all_data['Ranking Stream'].dropna().unique()
            selected_stream = st.selectbox('Select Ranking Stream', ranking_streams)
            states = all_data['State'].dropna().unique()
            selected_state = st.selectbox('Select State', states)
            if selected_stream and selected_state:
                stream_state_data = all_data[(all_data['Ranking Stream'] == selected_stream) & (all_data['State'] == selected_state)]
                st.write(stream_state_data[['College Name', 'Rank', 'City', 'State']].sort_values(by='Rank'))

                # Generate a simple summary paragraph
                state_summary = f"Overview of all colleges in {selected_state} for stream '{selected_stream}'."
                st.write(state_summary)

                # Display Graph for colleges in the selected state and stream
                watermark_text = "Your Watermark"
                graph_bytes = display_graph_for_college(stream_state_data, selected_state, watermark_text)

                # Create and Download Word Report
                word_file = create_word_report(f"{selected_state} - {selected_stream}", state_summary, stream_state_data[['College Name', 'Rank', 'City', 'State']], graph_bytes)
                
                # Add a download button for the Word document
                st.download_button(
                    label="Download Report as Word Document",
                    data=word_file,
                    file_name=f"{selected_state}_{selected_stream}_Ranking_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        elif ranking_type == 'All Colleges in a City':
            cities = all_data['City'].dropna().unique()
            selected_city = st.selectbox('Select City', cities)
            if selected_city:
                city_data = all_data[all_data['City'] == selected_city]
                st.write(city_data[['College Name', 'Rank', 'Ranking Stream', 'State']].sort_values(by='Rank'))
                st.write(f"Overview of all colleges in {selected_city}.")

                # Display Graph for colleges in the selected city
                watermark_text = "Your Watermark"
                graph_bytes = display_graph_for_college(city_data, selected_city, watermark_text)

                # Create and Download Word Report
                word_file = create_word_report(f"All Colleges in {selected_city}", f"Overview of all colleges in {selected_city}.", city_data[['College Name', 'Rank', 'Ranking Stream', 'State']], graph_bytes)
                
                # Add a download button for the Word document
                st.download_button(
                    label="Download Report as Word Document",
                    data=word_file,
                    file_name=f"All_Colleges_in_{selected_city}_Ranking_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        elif ranking_type == 'All Colleges in a State':
            states = all_data['State'].dropna().unique()
            selected_state = st.selectbox('Select State', states)
            if selected_state:
                state_data = all_data[all_data['State'] == selected_state]
                st.write(state_data[['College Name', 'Rank', 'Ranking Stream', 'City']].sort_values(by='Rank'))
                st.write(f"Overview of all colleges in {selected_state}.")

                # Display Graph for colleges in the selected state
                watermark_text = "collegedekho"
                graph_bytes = display_graph_for_college(state_data, selected_state, watermark_text)

                # Create and Download Word Report
                word_file = create_word_report(f"All Colleges in {selected_state}", f"Overview of all colleges in {selected_state}.", state_data[['College Name', 'Rank', 'Ranking Stream', 'City']], graph_bytes)
                
                # Add a download button for the Word document
                st.download_button(
                    label="Download Report as Word Document",
                    data=word_file,
                    file_name=f"All_Colleges_in_{selected_state}_Ranking_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
